"""
Document loaders for different file types.
"""
import logging
import os
from typing import List
from pathlib import Path

from langchain.schema import Document
from langchain.text_splitter import (
    RecursiveCharacterTextSplitter,
    CharacterTextSplitter,
)
from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
    UnstructuredMarkdownLoader,
)

from app.config import get_settings
from app.models import DocumentType, ChunkingStrategy

logger = logging.getLogger(__name__)


class DocumentLoaderFactory:
    """Factory for creating document loaders based on file type."""

    @staticmethod
    def load_document(file_path: str, file_type: DocumentType) -> List[Document]:
        """
        Load a document based on its type.

        Args:
            file_path: Path to the document
            file_type: Type of the document

        Returns:
            List of Document objects

        Raises:
            ValueError: If file type is not supported
        """
        logger.info(f"Loading document: {file_path} (type: {file_type})")

        try:
            if file_type == DocumentType.PDF:
                return DocumentLoaderFactory._load_pdf(file_path)
            elif file_type == DocumentType.TXT:
                return DocumentLoaderFactory._load_txt(file_path)
            elif file_type == DocumentType.DOCX:
                return DocumentLoaderFactory._load_docx(file_path)
            elif file_type in [DocumentType.MARKDOWN, DocumentType.MD]:
                return DocumentLoaderFactory._load_markdown(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            logger.error(f"Error loading document {file_path}: {e}")
            raise

    @staticmethod
    def _load_pdf(file_path: str) -> List[Document]:
        """Load PDF document."""
        loader = PyPDFLoader(file_path)
        documents = loader.load()

        # Add page numbers to metadata
        for i, doc in enumerate(documents):
            doc.metadata["page"] = i + 1
            doc.metadata["source"] = os.path.basename(file_path)

        logger.info(f"Loaded {len(documents)} pages from PDF")
        return documents

    @staticmethod
    def _load_txt(file_path: str) -> List[Document]:
        """Load text document."""
        loader = TextLoader(file_path, encoding="utf-8")
        documents = loader.load()

        # Add metadata
        for doc in documents:
            doc.metadata["source"] = os.path.basename(file_path)

        logger.info(f"Loaded text document with {len(documents)} parts")
        return documents

    @staticmethod
    def _load_docx(file_path: str) -> List[Document]:
        """Load DOCX document."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("python-docx is required for DOCX support. Install with: pip install python-docx")

        # Read DOCX content
        docx_doc = DocxDocument(file_path)

        # Extract text from paragraphs
        paragraphs = [para.text for para in docx_doc.paragraphs if para.text.strip()]
        content = "\n\n".join(paragraphs)

        # Create LangChain Document
        doc = Document(
            page_content=content,
            metadata={
                "source": os.path.basename(file_path),
                "paragraph_count": len(paragraphs)
            }
        )

        logger.info(f"Loaded DOCX with {len(paragraphs)} paragraphs")
        return [doc]

    @staticmethod
    def _load_markdown(file_path: str) -> List[Document]:
        """Load Markdown document."""
        try:
            # Try using UnstructuredMarkdownLoader
            loader = UnstructuredMarkdownLoader(file_path)
            documents = loader.load()
        except:
            # Fallback to reading as text
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            documents = [Document(
                page_content=content,
                metadata={"source": os.path.basename(file_path)}
            )]

        # Add metadata
        for doc in documents:
            doc.metadata["source"] = os.path.basename(file_path)

        logger.info(f"Loaded markdown document")
        return documents


class DocumentChunker:
    """Handle document chunking with different strategies."""

    def __init__(self, strategy: ChunkingStrategy = ChunkingStrategy.RECURSIVE):
        """Initialize chunker with strategy."""
        self.settings = get_settings()
        self.strategy = strategy

    def chunk_documents(
        self,
        documents: List[Document],
        chunk_size: int = None,
        chunk_overlap: int = None
    ) -> List[Document]:
        """
        Chunk documents using the configured strategy.

        Args:
            documents: List of documents to chunk
            chunk_size: Size of each chunk (uses config default if None)
            chunk_overlap: Overlap between chunks (uses config default if None)

        Returns:
            List of chunked documents
        """
        chunk_size = chunk_size or self.settings.chunk_size
        chunk_overlap = chunk_overlap or self.settings.chunk_overlap

        logger.info(f"Chunking {len(documents)} documents with strategy: {self.strategy}")

        if self.strategy == ChunkingStrategy.RECURSIVE:
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                length_function=len,
                separators=["\n\n", "\n", ". ", " ", ""]
            )
        elif self.strategy == ChunkingStrategy.CHARACTER:
            splitter = CharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                separator="\n"
            )
        else:
            # Default to recursive
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap
            )

        # Split documents
        chunks = splitter.split_documents(documents)

        # Add chunk index to metadata
        for i, chunk in enumerate(chunks):
            chunk.metadata["chunk_index"] = i

        logger.info(f"Created {len(chunks)} chunks")
        return chunks


def get_file_type_from_extension(filename: str) -> DocumentType:
    """
    Determine document type from filename extension.

    Args:
        filename: Name of the file

    Returns:
        DocumentType enum value

    Raises:
        ValueError: If file extension is not supported
    """
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return DocumentType.PDF
    elif ext == ".txt":
        return DocumentType.TXT
    elif ext == ".docx":
        return DocumentType.DOCX
    elif ext in [".md", ".markdown"]:
        return DocumentType.MARKDOWN
    else:
        raise ValueError(f"Unsupported file extension: {ext}")
