"""
Simplified document loader for conversational RAG system.
"""

import logging
from pathlib import Path
from typing import List
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.document_loaders import PyPDFLoader, TextLoader
from langchain_community.document_loaders import UnstructuredMarkdownLoader
import docx

logger = logging.getLogger(__name__)


def load_document(file_path: str) -> List[Document]:
    """
    Load a document from file path.

    Args:
        file_path: Path to the document

    Returns:
        List of Document objects
    """
    file_extension = Path(file_path).suffix.lower()

    try:
        if file_extension == ".pdf":
            loader = PyPDFLoader(file_path)
            documents = loader.load()
        elif file_extension == ".txt":
            loader = TextLoader(file_path)
            documents = loader.load()
        elif file_extension == ".docx":
            documents = load_docx(file_path)
        elif file_extension in [".md", ".markdown"]:
            loader = UnstructuredMarkdownLoader(file_path)
            documents = loader.load()
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")

        logger.info(f"Loaded {len(documents)} documents from {file_path}")
        return documents

    except Exception as e:
        logger.error(f"Error loading document {file_path}: {str(e)}")
        raise


def load_docx(file_path: str) -> List[Document]:
    """
    Load a DOCX document.

    Args:
        file_path: Path to DOCX file

    Returns:
        List with single Document object
    """
    doc = docx.Document(file_path)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

    return [
        Document(
            page_content=text,
            metadata={"source": Path(file_path).name},
        )
    ]


def chunk_documents(
    documents: List[Document],
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
) -> List[Document]:
    """
    Chunk documents using recursive character splitting.

    Args:
        documents: List of documents to chunk
        chunk_size: Size of each chunk
        chunk_overlap: Overlap between chunks

    Returns:
        List of chunked documents
    """
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
    )

    chunks = text_splitter.split_documents(documents)
    logger.info(f"Split {len(documents)} documents into {len(chunks)} chunks")

    return chunks
